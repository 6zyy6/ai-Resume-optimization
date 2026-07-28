from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import PurePath
from zipfile import BadZipFile, ZipFile

from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError
from pypdf.generic import DictionaryObject, IndirectObject


MAX_FILE_BYTES = 10 * 1024 * 1024
MAX_PDF_PAGES = 10
MAX_DOCX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 100
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@dataclass(frozen=True)
class ParsedResume:
    text: str
    page_count: int | None
    fallback_reason: str | None = None


class FileParseError(ValueError):
    def __init__(self, code: str, message: str) -> None:
        super().__init__(f"{code}: {message}")
        self.code = code
        self.message = message


def parse_resume_file(filename: str, mime: str, content: bytes) -> ParsedResume:
    if len(content) > MAX_FILE_BYTES:
        raise FileParseError("FILE_TOO_LARGE", "Files must not exceed 10 MiB")
    suffix = PurePath(filename).suffix.lower()
    if suffix == ".pdf":
        if mime != "application/pdf" or not content.startswith(b"%PDF-"):
            raise FileParseError("FILE_TYPE_UNSUPPORTED", "PDF signature does not match")
        return _parse_pdf(content)
    if suffix == ".docx":
        if mime != DOCX_MIME or not content.startswith(b"PK"):
            raise FileParseError("FILE_TYPE_UNSUPPORTED", "DOCX signature does not match")
        return _parse_docx(content)
    if suffix == ".txt":
        if mime not in {"text/plain", "text/plain; charset=utf-8"}:
            raise FileParseError("FILE_TYPE_UNSUPPORTED", "TXT MIME does not match")
        return _parse_txt(content)
    raise FileParseError("FILE_TYPE_UNSUPPORTED", "Only PDF, DOCX and TXT are accepted")


def _parse_pdf(content: bytes) -> ParsedResume:
    try:
        reader = PdfReader(BytesIO(content), strict=True)
        if reader.is_encrypted:
            raise FileParseError(
                "ENCRYPTED_PDF",
                "Encrypted PDF cannot be parsed; paste the resume text instead",
            )
        if len(reader.pages) > MAX_PDF_PAGES:
            raise FileParseError("PDF_PAGE_LIMIT", "PDF must not exceed 10 pages")
        _reject_pdf_actions(reader)
        text = "\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
    except FileParseError:
        raise
    except (PdfReadError, OSError, ValueError, KeyError) as error:
        raise FileParseError(
            "CORRUPT_FILE",
            "PDF is damaged; paste the resume text instead",
        ) from error
    if not text:
        raise FileParseError(
            "SCANNED_PDF",
            "No text layer found; paste the resume text instead",
        )
    return ParsedResume(text=text, page_count=len(reader.pages))


def _reject_pdf_actions(reader: PdfReader) -> None:
    root = reader.trailer.get("/Root")
    if root is None:
        return
    candidates = [
        root.get("/OpenAction"),
        root.get("/AA"),
        root.get("/AcroForm"),
        root.get("/Names"),
    ]
    for page in reader.pages:
        candidates.extend(
            [
                page.get("/AA"),
                page.get("/Annots"),
            ]
        )
    if any(_contains_pdf_javascript(candidate, set()) for candidate in candidates):
        raise FileParseError(
            "FILE_TYPE_UNSUPPORTED",
            "PDF embedded actions are not allowed",
        )


def _contains_pdf_javascript(value, seen: set[tuple[int, int]]) -> bool:
    if isinstance(value, IndirectObject):
        identity = (value.idnum, value.generation)
        if identity in seen:
            return False
        seen.add(identity)
        value = value.get_object()
    if isinstance(value, DictionaryObject):
        if value.get("/S") == "/JavaScript" or value.get("/JS") is not None:
            return True
        return any(
            _contains_pdf_javascript(child, seen)
            for child in value.values()
        )
    if isinstance(value, (list, tuple)):
        return any(_contains_pdf_javascript(child, seen) for child in value)
    return False


def _parse_docx(content: bytes) -> ParsedResume:
    try:
        with ZipFile(BytesIO(content)) as archive:
            entries = archive.infolist()
            if any(
                item.filename.lower().endswith("vbaproject.bin")
                or "activex" in item.filename.lower()
                for item in entries
            ):
                raise FileParseError("FILE_TYPE_UNSUPPORTED", "DOCX macros are not allowed")
            uncompressed = sum(item.file_size for item in entries)
            compressed = max(1, sum(item.compress_size for item in entries))
            if (
                uncompressed > MAX_DOCX_UNCOMPRESSED_BYTES
                or uncompressed / compressed > MAX_DOCX_COMPRESSION_RATIO
            ):
                raise FileParseError("FILE_TYPE_UNSUPPORTED", "Unsafe DOCX compression ratio")
        document = Document(BytesIO(content))
    except FileParseError:
        raise
    except (BadZipFile, KeyError, ValueError, OSError) as error:
        raise FileParseError(
            "CORRUPT_FILE",
            "DOCX is damaged; paste the resume text instead",
        ) from error
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            line = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if line:
                parts.append(line)
    if not parts:
        raise FileParseError("CORRUPT_FILE", "DOCX contains no readable text")
    return ParsedResume(text="\n".join(parts), page_count=None)


def _parse_txt(content: bytes) -> ParsedResume:
    try:
        text = content.decode("utf-8-sig").strip()
    except UnicodeDecodeError as error:
        raise FileParseError("FILE_PARSE_FAILED", "TXT must be UTF-8") from error
    if not text:
        raise FileParseError("FILE_PARSE_FAILED", "TXT contains no readable text")
    return ParsedResume(text=text, page_count=None)
