from io import BytesIO
import asyncio
import hashlib
from pathlib import Path
from datetime import datetime, timedelta, timezone
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from alembic import command
from alembic.config import Config
from sqlalchemy import create_engine, func, inspect, select, text
from sqlalchemy.exc import IntegrityError

from app.core.config import Settings
from app.db.models import (
    BulletFactLink,
    Fact,
    FactSource,
    File,
    IdempotencyRecord,
    Resume,
    ResumeImport,
    ResumeVersion,
    SourceRecord,
)
from app.modules.imports.parsers import FileParseError, parse_resume_file
from app.workers.execution import TaskExecutor, resolve_operation
from app.workers.pipeline import configure_pipeline_operations


def test_txt_import_rejects_magic_mime_and_size_mismatches():
    with pytest.raises(FileParseError, match="FILE_TYPE_UNSUPPORTED"):
        parse_resume_file("resume.pdf", "application/pdf", b"plain text")
    with pytest.raises(FileParseError, match="FILE_TOO_LARGE"):
        parse_resume_file("resume.txt", "text/plain", b"x" * (10 * 1024 * 1024 + 1))


def test_docx_parser_extracts_paragraphs_and_tables_and_rejects_macros():
    from docx import Document

    output = BytesIO()
    document = Document()
    document.add_paragraph("张三")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "项目经历"
    document.save(output)

    parsed = parse_resume_file(
        "resume.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        output.getvalue(),
    )
    assert "张三" in parsed.text
    assert "项目经历" in parsed.text

    macro = BytesIO()
    with ZipFile(macro, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("word/document.xml", "<document/>")
        archive.writestr("word/vbaProject.bin", b"macro")
    with pytest.raises(FileParseError, match="FILE_TYPE_UNSUPPORTED"):
        parse_resume_file(
            "resume.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            macro.getvalue(),
        )


def test_pdf_parser_reports_scanned_encrypted_corrupt_and_page_limit():
    from pypdf import PdfWriter

    scanned = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.write(scanned)
    with pytest.raises(FileParseError, match="SCANNED_PDF"):
        parse_resume_file("resume.pdf", "application/pdf", scanned.getvalue())

    encrypted = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.encrypt("secret")
    writer.write(encrypted)
    with pytest.raises(FileParseError, match="ENCRYPTED_PDF"):
        parse_resume_file("resume.pdf", "application/pdf", encrypted.getvalue())

    with pytest.raises(FileParseError, match="CORRUPT_FILE"):
        parse_resume_file("resume.pdf", "application/pdf", b"%PDF-1.7 broken")

    too_many = BytesIO()
    writer = PdfWriter()
    for _ in range(11):
        writer.add_blank_page(width=100, height=100)
    writer.write(too_many)
    with pytest.raises(FileParseError, match="PDF_PAGE_LIMIT"):
        parse_resume_file("resume.pdf", "application/pdf", too_many.getvalue())


def test_parsed_import_stays_draft_until_explicit_confirmation():
    from app.modules.imports.service import DraftImport

    draft = DraftImport.from_text("姓名：张三\n技能：Python")
    assert draft.confirmed is False
    assert draft.confirmed_facts == ()
    confirmed = draft.confirm()
    assert confirmed.confirmed is True
    assert len(confirmed.confirmed_facts) == 2


def test_0005_migration_creates_pipeline_columns(tmp_path, monkeypatch):
    database_path = tmp_path / "pipeline-migration.db"
    monkeypatch.setenv(
        "ALEMBIC_DATABASE_URL",
        f"sqlite+aiosqlite:///{database_path}",
    )
    api_root = Path(__file__).resolve().parents[1]
    config = Config(api_root / "alembic.ini")
    config.set_main_option("script_location", str(api_root / "migrations"))
    command.upgrade(config, "0005")
    engine = create_engine(f"sqlite:///{database_path}")
    try:
        schema = inspect(engine)
        assert "resume_imports" in schema.get_table_names()
        assert {"display_name", "status", "deleted_at"} <= {
            item["name"] for item in schema.get_columns("files")
        }
        assert {"original_text_encrypted", "reason", "risk_flags"} <= {
            item["name"] for item in schema.get_columns("suggestions")
        }
        assert {"download_name", "task_id"} <= {
            item["name"] for item in schema.get_columns("exports")
        }
    finally:
        engine.dispose()


def test_upload_import_and_confirmation_keeps_fact_gate(pipeline_client):
    client, sessions, storage = pipeline_client
    content = "技能：Python\n项目：校园服务平台".encode()
    digest = hashlib.sha256(content).hexdigest()
    token = client.post(
        "/v1/files/upload-tokens",
        json={
            "display_name": "resume.txt",
            "mime": "text/plain",
            "size": len(content),
            "sha256": digest,
            "purpose": "resume_import",
        },
        headers={"Idempotency-Key": "upload-import"},
    )
    assert token.status_code == 201
    file_id = token.json()["file_id"]
    file_row = asyncio.run(_file(sessions, file_id))
    uploaded = client.put(
        token.json()["upload_url"],
        content=content,
        headers={"Content-Type": "text/plain"},
    )
    assert uploaded.status_code == 204, uploaded.text

    confirmed = client.post(
        f"/v1/files/{file_id}/confirm-upload",
        headers={"Idempotency-Key": "confirm-file"},
    )
    created = client.post(
        "/v1/imports",
        json={"file_id": file_id},
        headers={"Idempotency-Key": "create-import"},
    )

    assert confirmed.status_code == 200
    assert created.status_code == 202
    assert created.json()["status"] == "queued"
    assert created.json()["task_id"]
    assert asyncio.run(_fact_count(sessions)) == 0
    configure_pipeline_operations(
        sessions,
        Settings(app_env="test", database_url="sqlite+aiosqlite://"),
        client.app.state.task_service,
        storage_override=storage,
    )
    task_result = asyncio.run(
        TaskExecutor(
            client.app.state.task_service,
            sleep=lambda _: None,
            jitter=lambda: 0,
        ).execute("usr_a", created.json()["task_id"], resolve_operation)
    )
    assert task_result["status"] == "succeeded"
    parsed = client.get(f"/v1/imports/{created.json()['id']}")
    assert parsed.json()["status"] == "parsed"

    final = client.post(
        f"/v1/imports/{created.json()['id']}/confirm",
        json={"title": "导入基础简历", "facts": []},
        headers={"Idempotency-Key": "confirm-import"},
    )
    replay = client.post(
        f"/v1/imports/{created.json()['id']}/confirm",
        json={"title": "导入基础简历", "facts": []},
        headers={"Idempotency-Key": "confirm-import"},
    )
    changed = client.post(
        f"/v1/imports/{created.json()['id']}/confirm",
        json={
            "title": "导入基础简历",
            "facts": [{"kind": "技能", "value": "Go"}],
        },
        headers={"Idempotency-Key": "confirm-import"},
    )
    assert final.status_code == 200
    assert final.json()["status"] == "confirmed"
    assert len(final.json()["fact_ids"]) == 2
    assert final.json()["resume_id"]
    assert final.json()["version_id"]
    assert replay.json() == final.json()
    assert changed.status_code == 409
    assert changed.json()["error"]["code"] == "IDEMPOTENCY_KEY_REUSED"
    assert asyncio.run(_fact_count(sessions)) == 2
    assert asyncio.run(_model_count(sessions, Resume)) == 1
    assert asyncio.run(_model_count(sessions, ResumeVersion)) == 1
    assert asyncio.run(_model_count(sessions, BulletFactLink)) == 2
    deleted = client.delete(
        f"/v1/files/{file_id}",
        headers={"Idempotency-Key": "delete-file"},
    )
    replayed_delete = client.delete(
        f"/v1/files/{file_id}",
        headers={"Idempotency-Key": "delete-file"},
    )
    assert deleted.status_code == 204
    assert replayed_delete.status_code == 204
    assert storage.get(file_row.object_key) is None


def test_upload_token_requires_and_enforces_idempotency(pipeline_client):
    """Retrying an upload-token request must not allocate a second object."""
    client, sessions, _ = pipeline_client
    content = b"resume"
    payload = {
        "display_name": "resume.txt",
        "mime": "text/plain",
        "size": len(content),
        "sha256": hashlib.sha256(content).hexdigest(),
        "purpose": "resume_import",
    }

    missing = client.post("/v1/files/upload-tokens", json=payload)
    first = client.post(
        "/v1/files/upload-tokens",
        json=payload,
        headers={"Idempotency-Key": "upload-token"},
    )
    replay = client.post(
        "/v1/files/upload-tokens",
        json=payload,
        headers={"Idempotency-Key": "upload-token"},
    )
    changed = client.post(
        "/v1/files/upload-tokens",
        json={
            **payload,
            "size": len(content) + 1,
            "sha256": hashlib.sha256(content + b"!").hexdigest(),
        },
        headers={"Idempotency-Key": "upload-token"},
    )

    assert missing.status_code == 422
    assert missing.json()["error"]["code"] == "IDEMPOTENCY_KEY_REQUIRED"
    assert first.status_code == 201
    assert replay.json() == first.json()
    assert changed.status_code == 409
    assert changed.json()["error"]["code"] == "IDEMPOTENCY_KEY_REUSED"
    assert asyncio.run(_file_count(sessions)) == 1


def test_import_finalize_rolls_back_every_resource_before_retry(pipeline_client):
    """A version failure must not leave facts, a resume, or a consumed retry key."""
    client, sessions, storage = pipeline_client
    import_id = _parsed_import(client, sessions, storage, "atomic")
    asyncio.run(
        _execute_sql(
            sessions,
            """
            CREATE TRIGGER fail_import_version
            BEFORE INSERT ON resume_versions
            BEGIN
              SELECT RAISE(ABORT, 'injected import finalize failure');
            END
            """,
        )
    )

    with pytest.raises(IntegrityError):
        client.post(
            f"/v1/imports/{import_id}/confirm",
            json={"title": "原子导入", "facts": []},
            headers={"Idempotency-Key": "atomic-finalize"},
        )

    assert asyncio.run(_import_status(sessions, import_id)) == "parsed"
    assert asyncio.run(_model_count(sessions, Fact)) == 0
    assert asyncio.run(_model_count(sessions, Resume)) == 0
    assert asyncio.run(_model_count(sessions, ResumeVersion)) == 0
    assert asyncio.run(_model_count(sessions, BulletFactLink)) == 0
    assert (
        asyncio.run(
            _idempotency_count(
                sessions,
                f"/v1/imports/{import_id}/confirm",
                "atomic-finalize",
            )
        )
        == 0
    )
    asyncio.run(_execute_sql(sessions, "DROP TRIGGER fail_import_version"))

    retried = client.post(
        f"/v1/imports/{import_id}/confirm",
        json={"title": "原子导入", "facts": []},
        headers={"Idempotency-Key": "atomic-finalize"},
    )

    assert retried.status_code == 200
    assert asyncio.run(_model_count(sessions, Fact)) == 2
    assert asyncio.run(_model_count(sessions, Resume)) == 1
    assert asyncio.run(_model_count(sessions, ResumeVersion)) == 1
    assert asyncio.run(_model_count(sessions, BulletFactLink)) == 2


def test_import_finalize_rejects_fabricated_fact_without_partial_resources(
    pipeline_client,
):
    client, sessions, storage = pipeline_client
    import_id = _parsed_import(client, sessions, storage, "fabricated")

    response = client.post(
        f"/v1/imports/{import_id}/confirm",
        json={
            "title": "无效导入",
            "facts": [{"kind": "skill", "value": "并不存在的 Go 经验"}],
        },
        headers={"Idempotency-Key": "fabricated-finalize"},
    )

    assert response.status_code == 422
    assert response.json()["error"]["code"] == "IMPORT_DRAFT_FACT_INVALID"
    assert asyncio.run(_import_status(sessions, import_id)) == "parsed"
    assert asyncio.run(_model_count(sessions, Fact)) == 0
    assert asyncio.run(_model_count(sessions, Resume)) == 0
    assert asyncio.run(_model_count(sessions, ResumeVersion)) == 0
    assert asyncio.run(_model_count(sessions, BulletFactLink)) == 0


def test_import_finalize_records_the_original_and_confirmed_edit_sources(
    pipeline_client,
):
    client, sessions, storage = pipeline_client
    import_id = _parsed_import(client, sessions, storage, "edited")

    response = client.post(
        f"/v1/imports/{import_id}/confirm",
        json={
            "title": "编辑后导入",
            "facts": [
                {
                    "kind": "skill",
                    "value": "熟练使用 Python",
                    "draft_index": 0,
                }
            ],
        },
        headers={"Idempotency-Key": "edited-finalize"},
    )

    assert response.status_code == 200
    assert asyncio.run(_model_count(sessions, Fact)) == 1
    assert asyncio.run(_model_count(sessions, SourceRecord)) == 2
    assert asyncio.run(_model_count(sessions, FactSource)) == 2
    link = asyncio.run(_first_bullet_link(sessions))
    assert len(link.fact_source_hashes_at_link) == 2


def test_upload_constraints_are_bound_to_the_signed_object(pipeline_client):
    client, sessions, storage = pipeline_client
    expected = b"safe text"
    token = client.post(
        "/v1/files/upload-tokens",
        json={
            "display_name": "resume.txt",
            "mime": "text/plain",
            "size": len(expected),
            "sha256": hashlib.sha256(expected).hexdigest(),
            "purpose": "resume_import",
        },
        headers={"Idempotency-Key": "upload-constraints"},
    )
    row = asyncio.run(_file(sessions, token.json()["file_id"]))
    storage.put(row.object_key, b"wrong txt", "text/plain")

    response = client.post(
        f"/v1/files/{row.id}/confirm-upload",
        headers={"Idempotency-Key": "mismatch"},
    )

    assert response.status_code == 422
    assert response.json()["error"]["code"] == "FILE_UPLOAD_MISMATCH"
    assert storage.get(row.object_key) is None


def test_expired_source_file_cleanup_removes_object_and_marks_file(pipeline_client):
    client, sessions, storage = pipeline_client
    content = b"resume"
    token = client.post(
        "/v1/files/upload-tokens",
        json={
            "display_name": "resume.txt",
            "mime": "text/plain",
            "size": len(content),
            "sha256": hashlib.sha256(content).hexdigest(),
            "purpose": "resume_import",
        },
        headers={"Idempotency-Key": "upload-cleanup"},
    )
    row = asyncio.run(_file(sessions, token.json()["file_id"]))
    storage.put(row.object_key, content, "text/plain")
    cutoff = datetime.now(timezone.utc)
    asyncio.run(_expire_file(sessions, row.id, cutoff - timedelta(seconds=1)))

    count = asyncio.run(
        client.app.state.import_service.cleanup_expired_files(cutoff)
    )
    cleaned = asyncio.run(_file_including_deleted(sessions, row.id))

    assert count == 1
    assert cleaned.status == "deleted"
    assert cleaned.deleted_at is not None
    assert storage.get(row.object_key) is None


async def _file(sessions, file_id):
    async with sessions() as session:
        return await session.get(File, file_id)


async def _fact_count(sessions):
    async with sessions() as session:
        return await session.scalar(select(func.count()).select_from(Fact))


async def _file_count(sessions):
    async with sessions() as session:
        return await session.scalar(select(func.count()).select_from(File))


async def _model_count(sessions, model):
    async with sessions() as session:
        return await session.scalar(select(func.count()).select_from(model))


async def _first_bullet_link(sessions):
    async with sessions() as session:
        return await session.scalar(select(BulletFactLink))


def _parsed_import(client, sessions, storage, key: str) -> str:
    content = "技能：Python\n项目：校园服务平台".encode()
    token = client.post(
        "/v1/files/upload-tokens",
        json={
            "display_name": "resume.txt",
            "mime": "text/plain",
            "size": len(content),
            "sha256": hashlib.sha256(content).hexdigest(),
            "purpose": "resume_import",
        },
        headers={"Idempotency-Key": f"{key}-upload"},
    )
    uploaded = client.put(
        token.json()["upload_url"],
        content=content,
        headers={"Content-Type": "text/plain"},
    )
    assert uploaded.status_code == 204
    file_id = token.json()["file_id"]
    assert client.post(
        f"/v1/files/{file_id}/confirm-upload",
        headers={"Idempotency-Key": f"{key}-confirm-upload"},
    ).status_code == 200
    created = client.post(
        "/v1/imports",
        json={"file_id": file_id},
        headers={"Idempotency-Key": f"{key}-create"},
    )
    configure_pipeline_operations(
        sessions,
        Settings(app_env="test", database_url="sqlite+aiosqlite://"),
        client.app.state.task_service,
        storage_override=storage,
    )
    task_result = asyncio.run(
        TaskExecutor(
            client.app.state.task_service,
            sleep=lambda _: None,
            jitter=lambda: 0,
        ).execute("usr_a", created.json()["task_id"], resolve_operation)
    )
    assert task_result["status"] == "succeeded"
    return created.json()["id"]


async def _execute_sql(sessions, statement: str) -> None:
    async with sessions.begin() as session:
        await session.execute(text(statement))


async def _import_status(sessions, import_id: str) -> str:
    async with sessions() as session:
        return await session.scalar(
            select(ResumeImport.status).where(ResumeImport.id == import_id)
        )


async def _idempotency_count(sessions, route: str, key: str) -> int:
    async with sessions() as session:
        return int(
            await session.scalar(
                select(func.count())
                .select_from(IdempotencyRecord)
                .where(
                    IdempotencyRecord.route == route,
                    IdempotencyRecord.key == key,
                )
            )
            or 0
        )


async def _expire_file(sessions, file_id, expires_at):
    async with sessions.begin() as session:
        row = await session.get(File, file_id)
        row.expires_at = expires_at


async def _file_including_deleted(sessions, file_id):
    async with sessions() as session:
        return await session.get(File, file_id)
